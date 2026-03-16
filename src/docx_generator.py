from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "override", "color": "#000000"}
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'left', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def json_to_docx(data, output_path):
    doc = Document()

    # Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    p_name = data.get('patient_information', {}).get('patient_name', 'Patient')
    title = doc.add_heading(f'VISIT DOCUMENTATION – {p_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Patient Information
    doc.add_heading('Patient Information', level=1)
    pt_info = data.get('patient_information', {})
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    cells = [
        ('Patient Name', pt_info.get('patient_name', '-'), 'DOB', pt_info.get('dob', '-')),
        ('Date of Service', pt_info.get('date_of_service', '-'), 'Physician', pt_info.get('physician', '-')),
        ('Scribe', pt_info.get('scribe', '-'), 'Facility', pt_info.get('facility', '-'))
    ]

    for i, row_data in enumerate(cells):
        row = table.rows[i].cells
        for j, val in enumerate(row_data):
            row[j].text = str(val)
            if j % 2 == 0:
                # Add bold to labels
                run = row[j].paragraphs[0].runs[0]
                run.bold = True

    doc.add_paragraph()

    # 2. Wound Assessment Table
    doc.add_heading('Wound Assessment Table', level=1)
    wounds = data.get('wounds', [])
    if wounds:
        # Field mapping
        ATTRS = [
            ("MIST Therapy", "mist_therapy"),
            ("Wound Location", "location"),
            ("Outcome", "outcome"),
            ("Wound Type", "type"),
            ("Wound Status", "status"),
            ("Measurements (L x W x D)", "measurements"),
            ("Area (sq cm)", "area_sq_cm"),
            ("Volume (cm³)", "volume_cu_cm"),
            ("Tunnels", "tunnels"),
            ("Max Depth (cm)", "max_depth"),
            ("Undermining (cm)", "undermining"),
            ("Stage / Grade", "stage_grade"),
            ("Exudate Amount", "exudate_amount"),
            ("Exudate Type", "exudate_type"),
            ("Odor", "odor"),
            ("Wound Margin", "wound_margin"),
            ("Periwound", "periwound"),
            ("Necrotic Material (%)", "necrotic_material"),
            ("Granulation (%)", "granulation"),
            ("Tissue Exposed", "tissue_exposed"),
            ("Debridement", "debridement"),
            ("Primary Dressing", "primary_dressing"),
            ("Secondary Dressing", "secondary_dressing"),
            ("Frequency", "frequency"),
            ("Special Equipment", "special_equipment")
        ]

        table = doc.add_table(rows=len(ATTRS) + 1, cols=len(wounds) + 1)
        table.style = 'Table Grid'
        
        # Header Row
        header_row = table.rows[0].cells
        header_row[0].text = "Field"
        header_row[0].paragraphs[0].runs[0].bold = True
        for i, w in enumerate(wounds):
            header_row[i+1].text = f"Wound {w.get('number', i+1)}"
            header_row[i+1].paragraphs[0].runs[0].bold = True
            
        # Data Rows
        for r_idx, (label, key) in enumerate(ATTRS):
            row = table.rows[r_idx+1].cells
            row[0].text = label
            row[0].paragraphs[0].runs[0].bold = True
            for c_idx, w in enumerate(wounds):
                row[c_idx+1].text = str(w.get(key, "-"))
    else:
        doc.add_paragraph("No wounds documented.")

    doc.add_paragraph()

    # 3. Detailed Visit Summaries
    doc.add_heading('Detailed Visit Summaries', level=1)
    for i, w in enumerate(wounds):
        num = w.get('number', str(i+1))
        wtype = w.get('type', '-')
        loc = w.get('location', '-')
        stage = w.get('stage_grade', '-')
        
        doc.add_heading(f'Wound {num}: {wtype}', level=2)
        
        p = doc.add_paragraph()
        p.add_run('Wound Location: ').bold = True
        p.add_run(str(loc))
        
        p = doc.add_paragraph()
        p.add_run('Stage: ').bold = True
        p.add_run(str(stage))
        
        doc.add_paragraph('Treatment:').runs[0].bold = True
        doc.add_paragraph(f'    • Primary Dressing: {w.get("primary_dressing", "-")}')
        doc.add_paragraph(f'    • Secondary Dressing: {w.get("secondary_dressing", "-")}')
        
        # Debridement checkboxes simulation
        d_sharp = "☑" if w.get("debridement_sharp") else "☐"
        d_none = "☑" if w.get("debridement_none") else "☐"
        doc.add_paragraph(f'    • Debridement:')
        doc.add_paragraph(f'        {d_sharp} Sharp debridement')
        doc.add_paragraph(f'        {d_none} No debridement')
        doc.add_paragraph(f'        Details: {w.get("debridement_details", "-")}')
        
        doc.add_paragraph(f'Offloading / Equipment: {w.get("offloading_equipment", "-")}').runs[0].bold = True
        
        clin_summ = w.get('clinical_summary', '')
        if clin_summ and clin_summ != "-":
            doc.add_paragraph('Summary:').runs[0].bold = True
            clean_summ = clin_summ.replace('\\n', '\n')
            clean_summ = re.sub(r'(?<=[.\s])\s*(Procedure:)', r'\n\1', clean_summ)
            clean_summ = re.sub(r'(?<=[.\s])\s*(Plan:)', r'\n\1', clean_summ)
            for line in clean_summ.split('\n'):
                if line.strip():
                    doc.add_paragraph(f'    {line.strip()}')

    doc.add_paragraph()

    # 4. Provider Comments
    doc.add_heading('Provider Comments', level=1)
    comments = data.get('comments', '-')
    if comments and comments != "-":
        clean_comm = comments.replace('\\n', '\n')
        for line in clean_comm.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
    else:
        doc.add_paragraph('______________')

    doc.add_paragraph()

    # 5. E/M Justification
    doc.add_heading('E/M Justification', level=1)
    em = data.get('em_justification', {})
    
    def _em_text(label, val):
        v = val if val and val != '-' else '______'
        return f'{label}: ___{v}___ minutes'

    doc.add_paragraph(_em_text('Time Spent Examining/Evaluating', em.get('time_spent_examining')))
    doc.add_paragraph(_em_text('Time Spent Documenting', em.get('time_spent_documenting')))
    doc.add_paragraph(_em_text('Time Spent Coordinating Care', em.get('time_spent_coordinating')))
    doc.add_paragraph(_em_text('Resolved Wound(s) - sign off', em.get('resolved_wound_sign_off')))
    doc.add_paragraph(f"Total Time: ___{em.get('total_time', '______')}___ minutes").runs[0].bold = True

    doc.save(output_path)
    return output_path
